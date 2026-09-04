"""生成 mock 简历测试文件（pdf/docx/图片），覆盖 file_parser 全部解析路径。

产出目录：backend/tests/fixtures/resume_samples/
  - mock_resume_backend.pdf    文本型 PDF（fitz 文本提取 + reflow 路径）
  - mock_resume_scanned.pdf    扫描件 PDF（无文本层，纯图 → OCR 兜底路径）
  - mock_resume_frontend.docx  Word（段落 + 技能表格 → docx 表格单元格提取路径）
  - mock_resume_data.png       图片简历（PaddleOCR 直接识别路径）
  - mock_resume_ops.jpg        图片简历 JPG 变体（大写扩展名兼容 + OCR）

版式按真实简历排版：居中姓名头 + 联系方式行 + 分隔线章节 + 时间线条目 + 项目符号。
姓名/电话/邮箱为**伪造**数据（号码段 138-0013-8xxx、邮箱 example.com），可顺带
验证 PII 脱敏链路；重跑本脚本会覆盖已生成文件。
"""

from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent / "tests" / "fixtures" / "resume_samples"
FONT_REG = "C:/Windows/Fonts/msyh.ttc"
FONT_BOLD = "C:/Windows/Fonts/msyhbd.ttc"

# 结构化简历数据：三种渲染器（PDF/Word/图片）共用
# 内容按 2026-09-04 图库实查的岗位技能画像对齐（Neo4j Position-[r]->Skill，weight 0.8 档全覆盖），
# 求职意向与图内岗位名一致以便推荐命中；0.4 档技能刻意不写，供差距/学习路径展示。
#   后端开发工程师: Python/Go/Java/Kubernetes/分布式技术
#   前端开发工程师: TypeScript/Vue.js/JavaScript/HTML5/React/CSS/Python
#   大数据开发工程师: SQL/Python/Apache Spark/Apache Kafka/ETL/数据管道/数据建模/K8s/CI-CD/分布式技术
#   运维工程师: Shell/Linux/Docker/Kubernetes/Python/SQL/Java/JavaScript
# blocks 类型：("entry", 左列, 中列, 右列) / ("bullet", 文本) / ("plain", 文本)
RESUMES = {
    "backend": {
        "name": "mock_resume_backend.pdf",
        "person": {
            "name": "张伟",
            "meta": "男 · 1999.06 · 本科 · 4 年经验 · 现居杭州",
            "phone": "138-0013-8000",
            "email": "zhangwei.dev@example.com",
        },
        "objective": "后端开发工程师",
        "blocks": [
            ("section", "教育背景"),
            ("entry", "2018.09 – 2022.06", "某理工大学", "软件工程 · 本科"),
            ("section", "工作经历"),
            ("entry", "2022.07 – 至今", "XX云计算科技有限公司", "后端开发工程师"),
            ("bullet", "基于 Python FastAPI 开发微服务网关与鉴权模块，PostgreSQL + Redis 支撑核心业务；"),
            ("bullet", "设计分布式事务与消息队列方案，QPS 峰值 3000 下 P99 稳定在 80ms 以内；"),
            ("bullet", "参与 Go 语言网关组件重写与 Java 存量服务维护，推动接口契约先行（OpenAPI）。"),
            ("section", "项目经验"),
            ("entry", "2023.05 – 2024.01", "统一鉴权中心", "核心开发"),
            ("bullet", "设计 OAuth2 + JWT 双令牌方案，支撑 12 个业务线接入，月活账号 40 万；"),
            ("bullet", "服务整体容器化迁移，基于 Kubernetes 编排与滚动发布，接入成本由 3 人日降至 0.5 人日。"),
            ("section", "专业技能"),
            ("bullet", "语言：Python（精通）、Go（熟练）、Java（熟练）；"),
            ("bullet", "架构：分布式技术（网关/消息队列/分布式事务）、Docker、Kubernetes；"),
            ("bullet", "存储：PostgreSQL、Redis。"),
        ],
    },
    "frontend": {
        "name": "mock_resume_frontend.docx",
        "person": {
            "name": "李娜",
            "meta": "女 · 2001.03 · 本科 · 3 年经验 · 现居上海",
            "phone": "138-0013-8001",
            "email": "lina.fe@example.com",
        },
        "objective": "前端开发工程师",
        "blocks": [
            ("section", "教育背景"),
            ("entry", "2019.09 – 2023.06", "某大学", "计算机科学与技术 · 本科"),
            ("section", "工作经历"),
            ("entry", "2023.07 – 至今", "XX互动娱乐有限公司", "前端开发工程师"),
            ("bullet", "负责运营活动页搭建系统与可视化大屏开发，基于 React 18 + TypeScript；"),
            ("bullet", "使用 Vue.js 3 维护另一条业务线，覆盖 JavaScript / HTML5 / CSS 全套 Web 技术；"),
            ("bullet", "编写 Python 构建脚本自动化资源打包，推动图表组件按需加载，首屏体积下降 35%。"),
            ("section", "项目经验"),
            ("entry", "2024.03 – 2024.11", "低代码活动搭建平台", "前端负责人"),
            ("bullet", "物料市场 + 拖拽画布方案，活动页产出周期由 3 天缩短至 2 小时；"),
            ("bullet", "设计物料协议与沙箱渲染层，接入 30+ 运营物料组件。"),
        ],
        # Word 特有：技能以表格呈现（覆盖 docx 表格单元格提取路径）
        "skills_table": [
            ("React / Redux", "精通"),
            ("TypeScript / JavaScript", "精通"),
            ("Vue.js", "熟练"),
            ("HTML5 / CSS", "精通"),
            ("ECharts / D3", "熟练"),
            ("Python（构建脚本）", "掌握"),
        ],
    },
    "data": {
        "name": "mock_resume_data.png",
        "person": {
            "name": "王强",
            "meta": "男 · 1998.11 · 本科 · 5 年经验 · 现居北京",
            "phone": "138-0013-8002",
            "email": "wangqiang.data@example.com",
        },
        "objective": "大数据开发工程师",
        "blocks": [
            ("section", "教育背景"),
            ("entry", "2017.09 – 2021.06", "某普通本科", "应用统计 · 本科"),
            ("section", "工作经历"),
            ("entry", "2021.07 – 至今", "XX数据科技有限公司", "大数据开发工程师"),
            ("bullet", "负责每日 2TB 增量日志的 ETL 数据管道，基于 SQL 与 Python 实现清洗入仓；"),
            ("bullet", "以 Apache Spark 替换单机批处理，Apache Kafka 构建实时链路，报表延迟由 24 小时降至 5 分钟；"),
            ("bullet", "完成核心域数据建模与表血缘治理，基于 Kubernetes 部署计算任务，CI/CD 流水线自动发布。"),
            ("section", "项目经验"),
            ("entry", "2023.02 – 2023.10", "实时数仓建设", "核心开发"),
            ("bullet", "Kafka + Spark Streaming 链路替换 T+1 批处理，数据及时率从 92% 提升至 99.5%；"),
            ("bullet", "Airflow 重构调度，任务失败自动重试与告警闭环。"),
            ("section", "专业技能"),
            ("bullet", "语言：SQL（精通）、Python（熟练）；"),
            ("bullet", "大数据：Apache Spark、Apache Kafka、ETL 数据管道、数据建模；"),
            ("bullet", "分布式技术：Spark 分布式计算、分区与数据倾斜治理；"),
            ("bullet", "平台：Kubernetes、CI/CD、Airflow。"),
        ],
    },
    "ops": {
        "name": "mock_resume_ops.JPG",
        "person": {
            "name": "陈静",
            "meta": "女 · 1997.09 · 本科 · 6 年经验 · 现居深圳",
            "phone": "138-0013-8003",
            "email": "chenjing.ops@example.com",
        },
        "objective": "运维工程师",
        "blocks": [
            ("section", "教育背景"),
            ("entry", "2016.09 – 2020.06", "某理工学院", "网络工程 · 本科"),
            ("section", "工作经历"),
            ("entry", "2020.07 – 至今", "XX互联网科技有限公司", "运维工程师"),
            ("bullet", "负责 200+ 节点 Kubernetes 集群与 Docker 容器化应用的运维，Linux 主机治理；"),
            ("bullet", "CI/CD 流水线治理，覆盖 Java 与 JavaScript（Node.js）应用的构建发布；"),
            ("bullet", "编写 Shell / Python 自动化巡检脚本，发布故障回滚时间由 30 分钟缩短至 3 分钟。"),
            ("section", "项目经验"),
            ("entry", "2024.01 – 2024.09", "多集群发布系统", "项目负责人"),
            ("bullet", "灰度发布 + 自动巡检，月均发布 400 次零重大事故；"),
            ("bullet", "SQL 慢查询排查与数据库巡检接入，Prometheus + Grafana 告警治理，误报率下降 60%。"),
            ("section", "专业技能"),
            ("bullet", "系统：Linux（精通）、Shell（精通）、Python（熟练）；"),
            ("bullet", "容器：Docker、Kubernetes（精通）；"),
            ("bullet", "其他：SQL（熟练）、Java（掌握）、JavaScript（掌握）。"),
        ],
    },
}

# 扫描件 PDF 复用后端简历：无文本层的纯图 PDF，走 _ocr_pdf 渲染+OCR 兜底
SCANNED_SOURCE_KEY = "backend"


# ---------------------------------------------------------------------------
# 图片渲染（PIL）—— 供图片简历与扫描件 PDF 共用
# ---------------------------------------------------------------------------

def _wrap(text: str, font, max_w: int) -> list[str]:
    """按像素宽度折行。"""
    from PIL import ImageFont

    lines, cur = [], ""
    for ch in text:
        if font.getlength(cur + ch) > max_w:
            lines.append(cur)
            cur = ch
        else:
            cur += ch
    if cur:
        lines.append(cur)
    return lines or [""]


def _draw_resume_image(person: dict, objective: str, blocks: list) -> "PIL.Image.Image":
    """按真实简历版式渲染白底图片：居中姓名头 + 章节分隔线 + 时间线条目。"""
    from PIL import Image, ImageDraw, ImageFont

    W = 1000
    M = 72                      # 页边距
    BODY_W = W - M * 2
    f_name = ImageFont.truetype(FONT_BOLD, 42)
    f_meta = ImageFont.truetype(FONT_REG, 22)
    f_section = ImageFont.truetype(FONT_BOLD, 26)
    f_body = ImageFont.truetype(FONT_REG, 22)
    line_h = 38
    section_gap = 22

    # 预计算高度
    height = 70 + 42 + 34 + 34 + line_h  # 姓名 + meta + 联系方式 + 求职意向 + 首个章节余量
    for kind, *rest in blocks:
        if kind == "section":
            height += line_h + section_gap
        elif kind == "entry":
            height += line_h
        else:
            text = rest[0] if kind == "bullet" else rest[0]
            prefix = "• " if kind == "bullet" else ""
            height += len(_wrap(prefix + text, f_body, BODY_W - 28)) * line_h
    height += 80

    img = Image.new("RGB", (W, height), "white")
    draw = ImageDraw.Draw(img)
    x, y = M, 64

    def center(text, font, dy):
        nonlocal y
        draw.text(((W - font.getlength(text)) / 2, y), text, font=font, fill="#111111")
        y += dy

    center(person["name"], f_name, 58)
    center(person["meta"], f_meta, 34)
    contact = f"电话：{person['phone']}    邮箱：{person['email']}"
    center(contact, f_meta, 40)

    # 求职意向行（居中加粗）+ 头部整宽分隔线
    center(f"求职意向：{objective}", f_section, 46)
    draw.line([(M, y), (W - M, y)], fill="#555555", width=2)
    y += 26

    for kind, *rest in blocks:
        if kind == "section":
            y += section_gap
            draw.text((x, y), rest[0], font=f_section, fill="#111111")
            y += line_h - 6
            draw.line([(x, y), (W - M, y)], fill="#999999", width=1)
            y += 12
        elif kind == "entry":
            left, mid, right = rest
            draw.text((x, y), f"{left}    {mid}", font=f_body, fill="#111111")
            rw = f_body.getlength(right)
            draw.text((W - M - rw, y), right, font=f_body, fill="#444444")
            y += line_h
        else:
            prefix = "• " if kind == "bullet" else ""
            indent = 28 if kind == "bullet" else 0
            for seg in _wrap(prefix + rest[0], f_body, BODY_W - indent):
                draw.text((x + indent, y), seg, font=f_body, fill="#222222")
                y += line_h
    return img


# ---------------------------------------------------------------------------
# PDF 渲染（fitz）—— 文本型简历
# ---------------------------------------------------------------------------

def _make_text_pdf(path: Path, person: dict, objective: str, blocks: list) -> None:
    import fitz

    f_reg = fitz.Font("china-s")
    f_bold = fitz.Font(fontfile=FONT_BOLD)  # 内置 CJK 无粗体，用微软雅黑粗体文件
    W, H, M = 595, 842, 56  # A4 pt
    doc = fitz.open()
    page = doc.new_page(width=W, height=H)

    def mount(p):
        p.insert_font(fontname="cjk", fontbuffer=f_reg.buffer)
        p.insert_font(fontname="cjkb", fontbuffer=f_bold.buffer)

    mount(page)
    y = 70

    def center(text, size, bold=False, color=(0.07, 0.07, 0.07)):
        nonlocal y
        f = f_bold if bold else f_reg
        w = f.text_length(text, fontsize=size)
        page.insert_text(((W - w) / 2, y), text, fontname="cjkb" if bold else "cjk",
                         fontsize=size, color=color)
        y += size + 12

    center(person["name"], 22, bold=True)
    center(person["meta"], 9.5, color=(0.3, 0.3, 0.3))
    contact = f"电话：{person['phone']}    邮箱：{person['email']}"
    center(contact, 9.5, color=(0.3, 0.3, 0.3))
    center(f"求职意向：{objective}", 12, bold=True)
    page.draw_line(fitz.Point(M, y - 4), fitz.Point(W - M, y - 4), color=(0.35, 0.35, 0.35), width=1)
    y += 20

    for kind, *rest in blocks:
        if y > H - 60:  # 简历单页放不下时换页
            page = doc.new_page(width=W, height=H)
            mount(page)
            y = 60
        if kind == "section":
            y += 10
            page.insert_text((M, y), rest[0], fontname="cjkb", fontsize=12)
            y += 6
            page.draw_line(fitz.Point(M, y), fitz.Point(W - M, y), color=(0.6, 0.6, 0.6), width=0.7)
            y += 16
        elif kind == "entry":
            left, mid, right = rest
            page.insert_text((M, y), f"{left}    {mid}", fontname="cjk", fontsize=10)
            rw = fitz.get_text_length(right, fontname="china-s", fontsize=10)
            page.insert_text((W - M - rw, y), right, fontname="cjk", fontsize=10,
                             color=(0.25, 0.25, 0.25))
            y += 17
        else:
            prefix = "• " if kind == "bullet" else ""
            indent = 14 if kind == "bullet" else 0
            avail = W - M * 2 - indent
            for seg in _wrap_pts(prefix + rest[0], avail):
                page.insert_text((M + indent, y), seg, fontname="cjk", fontsize=9.5,
                                 color=(0.13, 0.13, 0.13))
                y += 14.5

    doc.subset_fonts()  # 只嵌入用到的字形，全量 CJK 字体约 3.5MB → 子集后 KB 级
    doc.save(str(path))
    doc.close()


def _wrap_pts(text: str, max_w: float) -> list[str]:
    """按 pt 宽度折行（与正文 9.5pt 字号匹配）。"""
    import fitz

    lines, cur = [], ""
    for ch in text:
        if fitz.get_text_length(cur + ch, fontname="china-s", fontsize=9.5) > max_w:
            lines.append(cur)
            cur = ch
        else:
            cur += ch
    if cur:
        lines.append(cur)
    return lines or [""]


# ---------------------------------------------------------------------------
# Word 渲染（python-docx）
# ---------------------------------------------------------------------------

def _make_docx(path: Path, person: dict, objective: str, blocks: list,
               skills_table: list | None) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = docx.Document()

    def para(text, size=10.5, bold=False, align=None, color=None, space_after=4):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "微软雅黑"
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    para(person["name"], 22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(person["meta"], 9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55), space_after=2)
    para(f"电话：{person['phone']}    邮箱：{person['email']}", 9,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55), space_after=6)
    para(f"求职意向：{objective}", 13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    for kind, *rest in blocks:
        if kind == "section":
            para(rest[0], 13, bold=True, space_after=6)
        elif kind == "entry":
            left, mid, right = rest
            para(f"{left}    {mid}　　{right}", 10.5, bold=False, space_after=2)
        elif kind == "bullet":
            prefix = "• " if kind == "bullet" else ""
            para(f"{prefix}{rest[0]}", 10.5, space_after=2)

    if skills_table:
        para("专业技能", 13, bold=True, space_after=6)
        table = document.add_table(rows=len(skills_table), cols=2)
        table.style = "Table Grid"
        for i, (skill, level) in enumerate(skills_table):
            table.rows[i].cells[0].text = skill
            table.rows[i].cells[1].text = level
    document.save(str(path))


# ---------------------------------------------------------------------------
# 扫描件 PDF：图片整页插入，无文本层
# ---------------------------------------------------------------------------

def _make_scanned_pdf(path: Path, image) -> None:
    import fitz
    from io import BytesIO

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    buf = BytesIO()
    image.save(buf, format="PNG")
    page.insert_image(fitz.Rect(40, 40, 555, 800), stream=buf.getvalue())
    doc.save(str(path))
    doc.close()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    images: dict[str, object] = {}
    for key, spec in RESUMES.items():
        name = spec["name"]
        out = OUT_DIR / name
        person, objective, blocks = spec["person"], spec["objective"], spec["blocks"]
        image = _draw_resume_image(person, objective, blocks)
        images[key] = image
        if name.endswith(".pdf"):
            _make_text_pdf(out, person, objective, blocks)
        elif name.endswith(".docx"):
            _make_docx(out, person, objective, blocks, spec.get("skills_table"))
        else:
            if name.lower().endswith(".png"):
                out.write_bytes(_png_bytes(image))
            else:
                image.save(str(out), format="JPEG", quality=92)
        print(f"生成 {out.relative_to(OUT_DIR.parent.parent.parent)}")

    scanned = OUT_DIR / "mock_resume_scanned.pdf"
    _make_scanned_pdf(scanned, images[SCANNED_SOURCE_KEY])
    print(f"生成 {scanned.relative_to(OUT_DIR.parent.parent.parent)}")
    print(f"完成，共 {len(list(OUT_DIR.iterdir()))} 个文件 @ {OUT_DIR}")


def _png_bytes(image) -> bytes:
    from io import BytesIO

    buf = BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


if __name__ == "__main__":
    main()
