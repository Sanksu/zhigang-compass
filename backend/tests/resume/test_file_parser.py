"""简历文件文本抽取单元测试（AL-M3-04）。

覆盖 txt/docx 抽取、PDF 无文本（扫描件）抛错、文件缺失/类型不支持抛错。
"""

import pytest

from app.services.resume.file_parser import ResumeParseError, extract_text


class TestExtractText:
    def test_txt(self, tmp_path):
        path = tmp_path / "resume.txt"
        path.write_text("姓名：张三\n技能：Python、MySQL\n", encoding="utf-8")
        assert extract_text(path) == "姓名：张三\n技能：Python、MySQL"

    def test_docx_paragraphs_and_tables(self, tmp_path):
        import docx

        doc = docx.Document()
        doc.add_paragraph("姓名：李四")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "技能"
        table.rows[0].cells[1].text = "Go、Kubernetes"
        path = tmp_path / "resume.docx"
        doc.save(str(path))

        text = extract_text(path)
        assert "姓名：李四" in text
        assert "Go、Kubernetes" in text  # 表格内容也被抽取

    def test_pdf_without_text_raises(self, tmp_path, monkeypatch):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        path = tmp_path / "scan.pdf"
        with open(path, "wb") as f:
            writer.write(f)

        # 隔离真实 OCR：空白页场景下 OCR 引擎无文本可识别 → 抛"扫描件"错误。
        # 不加载 Paddle 模型（单测依赖外部模型下载会拖慢并受网络影响）。
        import app.services.resume.file_parser as fp

        monkeypatch.setattr(fp, "_ocr_bytes", lambda data, ocr: "")
        monkeypatch.setattr(
            fp, "_ocr_engine", lambda: object()
        )

        with pytest.raises(ResumeParseError, match="扫描件"):
            extract_text(path)

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(ResumeParseError, match="文件不存在"):
            extract_text(tmp_path / "nope.pdf")

    def test_pdf_chinese_text_layer_extracted(self, tmp_path):
        """回归（2026-08-09 P0 真实简历验证）：PyMuPDF 主提取中文 PDF。

        部分真实简历 PDF 中文用非 Unicode 自定义字体编码，pypdf 解码为乱码
        （实测同一 PDF pypdf 830 字乱码 vs fitz 1331 字正确），改用 fitz 主提取。
        用 fitz 生成含中文文本层的 PDF，验证 extract_text 正确提取。
        """
        import fitz

        path = tmp_path / "resume_cn.pdf"
        with fitz.open() as doc:
            page = doc.new_page()
            # china-s 为 PyMuPDF 内置 CJK 字体（插入中文必需，默认 Helvetica 不支持）
            page.insert_text((72, 72), "姓名：张三", fontname="china-s")
            page.insert_text((72, 100), "技能：Python、MySQL、React", fontname="china-s")
            doc.save(str(path))

        text = extract_text(path)
        assert "姓名：张三" in text
        assert "Python" in text
        assert "React" in text

    def test_ocr_image_extracts_text(self, tmp_path, monkeypatch):
        """图片简历走 OCR 成功路径：_ocr_bytes 收集 rec_texts 并拼接。"""
        from PIL import Image

        import app.services.resume.file_parser as fp

        img_path = tmp_path / "resume.png"
        Image.new("RGB", (100, 100), "white").save(img_path)

        class _FakeOcr:
            def predict(self, arr):
                return [{"rec_texts": ["姓名：张三", "技能：Python"]}]

        monkeypatch.setattr(fp, "_ocr_engine", lambda: _FakeOcr())

        text = extract_text(img_path)
        assert "姓名：张三" in text
        assert "Python" in text

    def test_ocr_image_empty_raises(self, tmp_path, monkeypatch):
        """OCR 无可识别文本 → 抛 ResumeParseError（不假成功返回）。"""
        from PIL import Image

        import app.services.resume.file_parser as fp

        img_path = tmp_path / "blank.png"
        Image.new("RGB", (100, 100), "white").save(img_path)

        class _EmptyOcr:
            def predict(self, arr):
                return [{"rec_texts": []}]

        monkeypatch.setattr(fp, "_ocr_engine", lambda: _EmptyOcr())

        with pytest.raises(ResumeParseError, match="无可识别文本"):
            extract_text(img_path)

    def test_unsupported_extension_raises(self, tmp_path):
        path = tmp_path / "resume.doc"  # 旧格式 .doc 明确不支持
        path.write_text("内容", encoding="utf-8")
        with pytest.raises(ResumeParseError, match="不支持的文件类型"):
            extract_text(path)
