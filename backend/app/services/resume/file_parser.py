"""简历文件文本抽取服务（AL-M3-04 + AL-M4-02 OCR）。

支持 pdf / docx / txt / 图片；PDF 无可提取文本（扫描件）或图片输入时
走 PaddleOCR 识别（ch_PP-OCRv4，懒加载）。OCR 不可用（模型下载失败/
依赖缺失）时抛 `ResumeParseError`，由调用方标记任务失败，不做假成功返回。

为 M4 简历解析（resume_parse）提供文本抽取前置能力，不含 LLM 抽取与 PII 脱敏。
"""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".bmp", ".webp"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}

# PDF 渲染 DPI：150 实测与 300 识别质量相当（2026-08-09 扫描件 OCR 速度评测），
# 且 300 会把文本行拆得更碎（4→8 行）、渲染耗时翻倍，故用 150 兼顾速度。
_OCR_RENDER_DPI = 150


class ResumeParseError(Exception):
    """简历文本抽取失败。"""


def extract_text(file_path: str | Path) -> str:
    """按扩展名抽取简历文本。

    Raises:
        ResumeParseError: 文件不存在 / 类型不支持 / 无可提取文本（含 OCR 失败）
    """
    path = Path(file_path)
    if not path.exists():
        raise ResumeParseError(f"文件不存在: {path}")

    ext = path.suffix.lower()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="replace").strip()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in _IMAGE_EXTS:
        return _ocr_image(path)
    raise ResumeParseError(
        f"不支持的文件类型: {ext}（仅支持 {'/'.join(sorted(SUPPORTED_EXTENSIONS))}，.doc 请转存为 .docx）"
    )


def _extract_pdf(path: Path) -> str:
    """PDF 文本提取：PyMuPDF 主提取 + reflow 双栏重组 + OCR 兜底。

    用 PyMuPDF（fitz）而非 pypdf 作主提取：pypdf 对部分真实简历 PDF 的中文
    （非 Unicode 自定义字体编码）解码为乱码，而 fitz 解码更强（实测同一 PDF
    pypdf 830 字乱码 vs fitz 1331 字正确，见 2026-08-09 P0 真实简历验证）。
    """
    import fitz  # PyMuPDF

    from app.services.resume.reflow import reflow_pdf

    with fitz.open(str(path)) as doc:
        text = "\n".join(page.get_text() or "" for page in doc).strip()
    if text:
        # 双栏版式（设计文档 §8.1）：pdfplumber 检测并按阅读顺序重组；
        # 重组为空串（单栏页）时保留 fitz 提取结果
        reflowed = reflow_pdf(path)
        return reflowed or text
    # 文本型提取为空 → 扫描件，走 OCR（pymupdf 渲染每页）
    return _ocr_pdf(path)


def _ocr_pdf(path: Path) -> str:
    """扫描件 PDF：pymupdf 渲染每页为图像 → PaddleOCR 识别。"""
    import fitz  # PyMuPDF

    ocr = _ocr_engine()
    texts = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            pix = page.get_pixmap(dpi=_OCR_RENDER_DPI)
            img_bytes = pix.tobytes("png")
            texts.append(_ocr_bytes(img_bytes, ocr))
    result = "\n".join(t for t in texts if t).strip()
    if not result:
        raise ResumeParseError(f"扫描件 OCR 无可识别文本: {path.name}")
    return result


def _ocr_image(path: Path) -> str:
    """图片简历：PaddleOCR 直接识别。"""
    result = _ocr_bytes(path.read_bytes(), _ocr_engine()).strip()
    if not result:
        raise ResumeParseError(f"图片 OCR 无可识别文本: {path.name}")
    return result


_ocr = None


def _ocr_engine():
    """PaddleOCR 懒加载单例（首次调用下载 PP-OCRv6 模型到用户目录）。

    通过 paddlex `create_pipeline` 构建并直传 `engine_config`（run_mode=paddle）：
    paddle 3.3.x 默认 CPU 走 PIR 新执行器 + oneDNN，触发 onednn_instruction 的
    ConvertPirAttribute2RuntimeAttribute 未实现错误，run_mode=paddle 可禁用 oneDNN 规避。
    """
    global _ocr
    if _ocr is None:
        try:
            from paddlex.inference import load_pipeline_config
            from paddlex.inference.pipelines import create_pipeline
        except ImportError as e:  # 依赖未安装
            raise ResumeParseError(f"OCR 依赖缺失: {e}") from e

        config = load_pipeline_config("OCR")
        # 简历正文通常正向排版，禁用文档方向/解卷/文本行方向分类，减少 3 个子模型加载
        config["SubPipelines"]["DocPreprocessor"]["use_doc_orientation_classify"] = False
        config["SubPipelines"]["DocPreprocessor"]["use_doc_unwarping"] = False
        config["use_textline_orientation"] = False
        for name in ("DocOrientationClassify", "DocUnwarping"):
            config["SubPipelines"]["DocPreprocessor"]["SubModules"].pop(name, None)
        _ocr = create_pipeline(
            config=config,
            device="cpu",
            engine_config={"run_mode": "paddle", "device_type": "cpu"},
        )
    return _ocr


def _ocr_bytes(data: bytes, ocr) -> str:
    """OCR 识别图像字节（paddlex predict 返回含 rec_texts 的结果字典列表）。"""
    import numpy as np
    from PIL import Image
    from io import BytesIO

    image = Image.open(BytesIO(data)).convert("RGB")
    results = ocr.predict(np.array(image))
    lines = []
    for res in results:
        lines.extend(res.get("rec_texts", []) or [])
    return "\n".join(lines)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # 表格单元格也是简历正文（如技能清单/经历表）
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeParseError(f"DOCX 无可提取文本: {path.name}")
    return text
